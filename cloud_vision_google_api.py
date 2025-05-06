import os
import base64
import json
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from google.oauth2 import service_account
from google.auth.transport.requests import Request

# مسیرها
SERVICE_ACCOUNT_FILE = "ocr-oiec-b2bd50d42cba.json"
INPUT_FOLDER = "images"
OUTPUT_FOLDER = "outputs"
os.makedirs(f"{OUTPUT_FOLDER}/txt", exist_ok=True)
os.makedirs(f"{OUTPUT_FOLDER}/json", exist_ok=True)
os.makedirs(f"{OUTPUT_FOLDER}/docx", exist_ok=True)

# گرفتن access token
credentials = service_account.Credentials.from_service_account_file(
    SERVICE_ACCOUNT_FILE,
    scopes=["https://www.googleapis.com/auth/cloud-platform"],
)
credentials.refresh(Request())
access_token = credentials.token

# تنظیمات درخواست API
API_URL = "https://vision.googleapis.com/v1/images:annotate"
HEADERS = {"Authorization": f"Bearer {access_token}"}

# تنظیم فونت فارسی در Word
def style_paragraph(paragraph):
    run = paragraph.runs[0]
    run.font.name = 'B Nazanin'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# پیمایش فایل‌ها
for file in os.listdir(INPUT_FOLDER):
    if not file.lower().endswith(('.png', '.jpg', '.jpeg')):
        continue

    image_path = os.path.join(INPUT_FOLDER, file)
    print(f"در حال پردازش: {file}")

    with open(image_path, "rb") as image_file:
        encoded_image = base64.b64encode(image_file.read()).decode("utf-8")

    body = {
        "requests": [
            {
                "image": {"content": encoded_image},
                "features": [{"type": "DOCUMENT_TEXT_DETECTION"}],  # یا TEXT_DETECTION
                "imageContext": {"languageHints": ["fa", "en"]}
            }
        ]
    }

    response = requests.post(API_URL, headers=HEADERS, json=body)

    if response.status_code != 200:
        print(f"❌ خطا در پردازش {file}: {response.status_code}")
        continue

    result = response.json()
    response_data = result.get("responses", [{}])[0]

    # 🔸 ذخیره JSON
    json_path = f"{OUTPUT_FOLDER}/json/{os.path.splitext(file)[0]}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(response_data, f, ensure_ascii=False, indent=2)

    # 🔸 استخراج متن ساده
    full_text = response_data.get("fullTextAnnotation", {}).get("text", "")

    # ذخیره متن در txt
    txt_path = f"{OUTPUT_FOLDER}/txt/{os.path.splitext(file)[0]}.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(full_text)

    # 🔸 ساخت فایل Word
    doc = Document()
    doc.add_heading(f"OCR result for {file}", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    for page in response_data.get("fullTextAnnotation", {}).get("pages", []):
        for block in page.get("blocks", []):
            for paragraph in block.get("paragraphs", []):
                para_text = ""
                for word in paragraph.get("words", []):
                    for symbol in word.get("symbols", []):
                        para_text += symbol.get("text", "")
                    para_text += " "
                p = doc.add_paragraph(para_text.strip())
                style_paragraph(p)

    doc_path = f"{OUTPUT_FOLDER}/docx/{os.path.splitext(file)[0]}.docx"
    doc.save(doc_path)

print("✅ تمام تصاویر با موفقیت OCR و ذخیره شدند.")
