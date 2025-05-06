# ocr_helper.py

import os
import base64
import json
import requests
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from google.oauth2 import service_account
from google.auth.transport.requests import Request

class OCRHelper:
    def __init__(self, service_account_path, detection_type="TEXT_DETECTION"):
        self.service_account_path = service_account_path
        self.detection_type = detection_type
        self.access_token = self.get_access_token()

    def get_access_token(self):
        credentials = service_account.Credentials.from_service_account_file(
            self.service_account_path,
            scopes=["https://www.googleapis.com/auth/cloud-platform"],
        )
        credentials.refresh(Request())
        return credentials.token

    def style_paragraph(self, paragraph):
        """Apply font, size, alignment and BiDi to the paragraph."""
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.name = 'B Nazanin'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            run.font.size = Pt(14)

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def fix_parentheses_for_rtl(self, word):
        """Properly adjust parentheses for Persian (RTL) text inside brackets."""
        RLM = '\u200F'

        contains_farsi = any('\u0600' <= c <= '\u06FF' for c in word)

        if contains_farsi:
            word = word.replace('(', RLM + '(' + RLM)
            word = word.replace(')', RLM + ')' + RLM)

        return word



    def set_run_rtl(self, run):
        """Force a run to be Right-to-Left."""
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)

    def set_run_ltr(self, run):
        """Force a run to be Left-to-Right."""
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '0')
        rPr.append(rtl)

    def is_ltr_word(self, word):
        """Determine if a word should be LTR based on English letters, numbers, or symbols."""
        for c in word:
            if (
                'A' <= c <= 'Z' or
                'a' <= c <= 'z' or
                '0' <= c <= '9' or
                c in ['/', '\\', '-', '_', ':', ';', '.', ',', '[', ']', '{', '}', '@', '#', '$', '%', '^', '&', '*', '+', '=', '~', '`', '<', '>', '|', '"', "'"]
            ):
                return True
        return False

    def add_mixed_text(self, paragraph, text):
        """Add Persian-English mixed text with proper direction word by word."""
        words = text.split(' ')
        for word in words:
            run = paragraph.add_run()
            if self.is_ltr_word(word):
                run.add_text(word + ' ')
                self.set_run_ltr(run)
            else:
                fixed_word = self.fix_parentheses_for_rtl(word)
                run.add_text(fixed_word + ' ')
                self.set_run_rtl(run)


    def ocr_image(self, image_path, save_txt=True, save_json=True, save_docx=True):
        """Perform OCR on an image and save outputs."""
        base_filename = os.path.splitext(image_path)[0]

        try:
            with open(image_path, "rb") as image_file:
                encoded_image = base64.b64encode(image_file.read()).decode("utf-8")
        except Exception as e:
            raise Exception(f"Failed to read image {image_path}: {e}")

        url = "https://vision.googleapis.com/v1/images:annotate"
        headers = {"Authorization": f"Bearer {self.access_token}"}
        body = {
            "requests": [
                {
                    "image": {"content": encoded_image},
                    "features": [{"type": self.detection_type}],
                    "imageContext": {"languageHints": ["fa", "en"]}
                }
            ]
        }

        try:
            response = requests.post(url, headers=headers, json=body)
            response.raise_for_status()
            result = response.json()
            response_data = result.get("responses", [{}])[0]
        except Exception as e:
            raise Exception(f"Vision API error for {image_path}: {e}")

        # Save outputs
        if save_json:
            with open(f"{base_filename}.json", "w", encoding="utf-8") as f:
                json.dump(response_data, f, ensure_ascii=False, indent=2)

        full_text = response_data.get("fullTextAnnotation", {}).get("text", "")
        if save_txt:
            with open(f"{base_filename}.txt", "w", encoding="utf-8") as f:
                f.write(full_text)

        if save_docx:
            doc = Document()
            for page in response_data.get("fullTextAnnotation", {}).get("pages", []):
                for block in page.get("blocks", []):
                    for paragraph in block.get("paragraphs", []):
                        para_text = ""
                        for word in paragraph.get("words", []):
                            for symbol in word.get("symbols", []):
                                para_text += symbol.get("text", "")
                            para_text += " "
                        if para_text.strip():
                            p = doc.add_paragraph()
                            self.add_mixed_text(p, para_text.strip())
                            self.style_paragraph(p)

            doc.save(f"{base_filename}.docx")
